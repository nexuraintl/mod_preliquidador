"""
EXTRACTOR DE DOCUMENTOS CON GUARDADO AUTOMÁTICO - VERSIÓN CORREGIDA
==================================================================

Módulo para extraer texto de diferentes tipos de archivos:
- PDFs (con conversión a imagen para OCR fallback)
- Imágenes (JPG, PNG) usando Google Vision OCR
- Excel (XLSX, XLS)
- Word (DOCX, DOC)

CORRECCIÓN: PDF → Imagen → OCR para casos de poco texto extraído.

Autor: Miguel Angel Jaramillo Durango
"""

import os
import io
import json
import logging
from pathlib import Path
from typing import Dict, Any
from datetime import datetime

# Procesamiento de archivos
import PyPDF2
from PIL import Image
import pandas as pd
from docx import Document

# Nuevas dependencias para PDF → Imagen
try:
    import pdf2image
    PDF2IMAGE_DISPONIBLE = True
except ImportError:
    PDF2IMAGE_DISPONIBLE = False
    logging.warning("pdf2image no disponible - OCR fallback para PDF limitado")

try:
    import fitz  # PyMuPDF
    PYMUPDF_DISPONIBLE = True
except ImportError:
    PYMUPDF_DISPONIBLE = False

# Google Vision para OCR
from google.cloud import vision

# FastAPI
from fastapi import UploadFile

# Configuración de logging
logger = logging.getLogger(__name__)

# ===============================
# PROCESADOR DE ARCHIVOS CON GUARDADO AUTOMÁTICO
# ===============================

class ProcesadorArchivos:
    """
    Extrae texto de diferentes tipos de archivos usando las mejores técnicas
    disponibles para cada formato.
    
    CORRECCIÓN: Maneja correctamente PDF → Imagen → OCR
    """
    
    def __init__(self):
        """Inicializa el procesador con configuración de OCR y carpetas de guardado"""
        self.vision_client = self._configurar_vision()
        self._crear_carpetas_guardado()
        self._verificar_dependencias_pdf()
        logger.info("ProcesadorArchivos inicializado con guardado automático")
    
    def _verificar_dependencias_pdf(self):
        """Verifica y reporta las dependencias disponibles para PDF → Imagen"""
        if PDF2IMAGE_DISPONIBLE:
            logger.info("✅ pdf2image disponible para conversión PDF → Imagen")
        elif PYMUPDF_DISPONIBLE:
            logger.info("✅ PyMuPDF disponible para conversión PDF → Imagen")
        else:
            logger.warning("⚠️ Sin dependencias para PDF → Imagen. OCR fallback limitado")
            logger.warning("   Instala: pip install pdf2image PyMuPDF")
    
    def _configurar_vision(self):
        """
        Configura Google Vision para OCR (opcional).
        
        Returns:
            vision.ImageAnnotatorClient o None si no está configurado
        """
        try:
            from dotenv import load_dotenv
            load_dotenv()
            
            credentials_path = os.getenv("GOOGLE_APPLICATION_CREDENTIALS")
            if credentials_path and os.path.exists(credentials_path):
                vision_client = vision.ImageAnnotatorClient()
                logger.info("Google Vision configurado correctamente para OCR")
                return vision_client
            else:
                logger.warning("Google Vision no configurado - OCR limitado")
                return None
                
        except Exception as e:
            logger.warning(f"No se pudo configurar Google Vision: {e}")
            return None
    
    def _crear_carpetas_guardado(self):
        """Crea las carpetas necesarias para guardar extracciones"""
        try:
            # Crear carpeta base de extracciones
            self.carpeta_base = Path("Results/Extracciones")
            self.carpeta_base.mkdir(parents=True, exist_ok=True)
            
            # Crear carpeta por fecha actual
            fecha_hoy = datetime.now().strftime("%Y-%m-%d")
            self.carpeta_fecha = self.carpeta_base / fecha_hoy
            self.carpeta_fecha.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Carpetas de guardado creadas: {self.carpeta_fecha}")
            
        except Exception as e:
            logger.error(f"Error creando carpetas de guardado: {e}")
            # Fallback a carpeta actual
            self.carpeta_fecha = Path(".")
    
    def _guardar_texto_extraido(self, nombre_archivo: str, texto_extraido: str, 
                               metodo_extraccion: str, metadatos: Dict = None) -> str:
        """
        Guarda el texto extraído en un archivo organizado.
        
        Args:
            nombre_archivo: Nombre del archivo original
            texto_extraido: Texto que se extrajo
            metodo_extraccion: Método usado (PDF, OCR, EXCEL, WORD)
            metadatos: Información adicional sobre la extracción
            
        Returns:
            str: Ruta del archivo guardado
        """
        try:
            # Crear timestamp único
            timestamp = datetime.now().strftime("%H%M%S")
            
            # Limpiar nombre de archivo para usar como nombre base
            nombre_base = "".join(c for c in nombre_archivo if c.isalnum() or c in "._-")
            if len(nombre_base) > 50:  # Limitar longitud
                nombre_base = nombre_base[:50]
            
            # Crear nombre de archivo único
            nombre_salida = f"{timestamp}_{metodo_extraccion}_{nombre_base}.txt"
            ruta_archivo = self.carpeta_fecha / nombre_salida
            
            # Preparar contenido completo
            contenido_completo = f"""EXTRACCIÓN DE TEXTO - PRELIQUIDADOR v2.0
============================================

INFORMACIÓN DEL ARCHIVO:
- Archivo original: {nombre_archivo}
- Método de extracción: {metodo_extraccion}
- Fecha y hora: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
- Caracteres extraídos: {len(texto_extraido)}

METADATOS ADICIONALES:
{json.dumps(metadatos or {}, indent=2, ensure_ascii=False)}

============================================
TEXTO EXTRAÍDO:
============================================

{texto_extraido}

============================================
FIN DE LA EXTRACCIÓN
============================================
"""
            
            # Guardar archivo
            with open(ruta_archivo, 'w', encoding='utf-8') as f:
                f.write(contenido_completo)
            
            logger.info(f"✅ Texto extraído guardado: {ruta_archivo}")
            logger.info(f"📊 Caracteres extraídos: {len(texto_extraido)}")
            
            return str(ruta_archivo)
            
        except Exception as e:
            logger.error(f"❌ Error guardando texto extraído: {e}")
            return f"Error guardando: {str(e)}"
    
    def _validar_pdf(self, contenido_pdf: bytes, nombre_archivo: str) -> Dict[str, Any]:
        """
        Valida un PDF antes de intentar conversión a imagen.
        
        Args:
            contenido_pdf: Contenido binario del PDF
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            Dict con información de validación
        """
        validacion = {
            "valido": False,
            "error": None,
            "info": {},
            "metodo_recomendado": None
        }
        
        try:
            # Verificar tamaño mínimo
            if len(contenido_pdf) < 100:
                validacion["error"] = f"PDF demasiado pequeño: {len(contenido_pdf)} bytes"
                return validacion
            
            # Verificar header PDF
            if not contenido_pdf.startswith(b'%PDF-'):
                validacion["error"] = "No es un archivo PDF válido (header incorrecto)"
                return validacion
            
            # Intentar abrir con PyPDF2 para validación básica
            try:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(contenido_pdf))
                num_paginas = len(pdf_reader.pages)
                
                validacion["info"]["paginas"] = num_paginas
                validacion["info"]["tamaño_bytes"] = len(contenido_pdf)
                
                if num_paginas == 0:
                    validacion["error"] = "PDF sin páginas"
                    return validacion
                
                # Verificar si está protegido
                if pdf_reader.is_encrypted:
                    validacion["error"] = "PDF protegido con contraseña"
                    return validacion
                
                validacion["valido"] = True
                
                # Recomendar método según dependencias disponibles
                if PDF2IMAGE_DISPONIBLE:
                    validacion["metodo_recomendado"] = "pdf2image"
                elif PYMUPDF_DISPONIBLE:
                    validacion["metodo_recomendado"] = "PyMuPDF"
                else:
                    validacion["error"] = "No hay librerías de conversión instaladas"
                    validacion["valido"] = False
                
            except Exception as e:
                validacion["error"] = f"PDF corrupto o inválido: {str(e)}"
                return validacion
                
        except Exception as e:
            validacion["error"] = f"Error validando PDF: {str(e)}"
        
        return validacion
    
    def _convertir_pdf_a_imagenes(self, contenido_pdf: bytes, nombre_archivo: str) -> list:
        """
        Convierte páginas de PDF a imágenes para OCR.
        
        Args:
            contenido_pdf: Contenido binario del PDF
            nombre_archivo: Nombre del archivo para logging
            
        Returns:
            list: Lista de imágenes en bytes (JPEG)
        """
        imagenes = []
        
        try:
            # Validar que el contenido no esté vacío
            if not contenido_pdf or len(contenido_pdf) < 100:
                logger.error(f"❌ PDF demasiado pequeño o vacío: {len(contenido_pdf)} bytes")
                return []
            
            # Intentar con pdf2image primero
            if PDF2IMAGE_DISPONIBLE:
                try:
                    from pdf2image import convert_from_bytes
                    
                    logger.info(f"🔄 Convirtiendo PDF a imágenes con pdf2image: {nombre_archivo}")
                    
                    # Convertir PDF a imágenes con configuración robusta
                    pages = convert_from_bytes(
                        contenido_pdf,
                        dpi=300,  # Alta calidad para OCR
                        fmt='JPEG',
                        thread_count=1,  # Evitar problemas de concurrencia
                        first_page=1,
                        last_page=100  # Limitar a 100 páginas máximo
                    )
                    
                    # Convertir cada página a bytes
                    for i, page in enumerate(pages):
                        img_byte_arr = io.BytesIO()
                        page.save(img_byte_arr, format='JPEG', quality=95)
                        imagenes.append(img_byte_arr.getvalue())
                    
                    logger.info(f"✅ pdf2image: {len(imagenes)} páginas convertidas")
                    return imagenes
                    
                except Exception as e:
                    logger.error(f"❌ Error con pdf2image: {e}")
                    # Continuar con PyMuPDF
            
            # Intentar con PyMuPDF como alternativa
            if PYMUPDF_DISPONIBLE:
                try:
                    logger.info(f"🔄 Convirtiendo PDF a imágenes con PyMuPDF: {nombre_archivo}")
                    
                    # Abrir PDF desde bytes
                    pdf_document = fitz.open(stream=contenido_pdf, filetype="pdf")
                    
                    # Verificar que el PDF se abrió correctamente
                    if pdf_document.page_count == 0:
                        logger.error(f"❌ PDF sin páginas válidas")
                        pdf_document.close()
                        return []
                    
                    # Limitar número de páginas
                    max_pages = min(pdf_document.page_count, 100)
                    logger.info(f"📄 Procesando {max_pages} de {pdf_document.page_count} páginas")
                    
                    # Convertir cada página
                    for page_num in range(max_pages):
                        page = pdf_document.load_page(page_num)
                        
                        # Renderizar página como imagen
                        mat = fitz.Matrix(3.0, 3.0)  # Factor de escala para calidad
                        pix = page.get_pixmap(matrix=mat)
                        
                        # Convertir a bytes JPEG
                        img_data = pix.tobytes("jpeg")
                        imagenes.append(img_data)
                    
                    pdf_document.close()
                    logger.info(f"✅ PyMuPDF: {len(imagenes)} páginas convertidas")
                    return imagenes
                    
                except Exception as e:
                    logger.error(f"❌ Error con PyMuPDF: {e}")
            
            # Si llegamos aquí, ambos métodos fallaron
            logger.error(f"❌ Todos los métodos de conversión fallaron")
            logger.error(f"   pdf2image disponible: {PDF2IMAGE_DISPONIBLE}")
            logger.error(f"   PyMuPDF disponible: {PYMUPDF_DISPONIBLE}")
            
            if not PDF2IMAGE_DISPONIBLE and not PYMUPDF_DISPONIBLE:
                logger.error(f"   CAUSA: No hay librerías de conversión instaladas")
                logger.error(f"   SOLUCIÓN: pip install pdf2image PyMuPDF")
            
            return []
                
        except Exception as e:
            logger.error(f"❌ Error general convirtiendo PDF a imágenes: {e}")
            logger.error(f"   Archivo: {nombre_archivo}")
            logger.error(f"   Tamaño: {len(contenido_pdf)} bytes")
            return []
    
    async def procesar_archivo(self, archivo: UploadFile) -> str:
        """
        Procesa un archivo y extrae su texto usando el método más apropiado.
        GUARDA AUTOMÁTICAMENTE el texto extraído para auditoría.
        
        Args:
            archivo: Archivo subido via FastAPI
            
        Returns:
            str: Texto extraído del archivo
            
        Raises:
            ValueError: Si el tipo de archivo no es soportado
        """
        if not archivo.filename:
            raise ValueError("Archivo sin nombre")
        
        extension = Path(archivo.filename).suffix.lower()
        contenido = await archivo.read()
        
        logger.info(f"📄 Procesando archivo: {archivo.filename} ({extension})")
        
        # Determinar método de extracción según extensión
        if extension == '.pdf':
            texto = await self.extraer_texto_pdf(contenido, archivo.filename)
            # Si se extrajo muy poco texto de PDF, intentar OCR
            if len(texto.strip()) < 1000 and not texto.startswith("Error"):
                logger.info("🔄 Poco texto extraído de PDF, intentando OCR con conversión a imagen...")
                texto_ocr = await self.extraer_texto_pdf_con_ocr(contenido, archivo.filename)
                if texto_ocr and len(texto_ocr) > len(texto) and not texto_ocr.startswith("Error"):
                    logger.info("✅ OCR proporcionó mejor resultado que extracción de PDF")
                    return texto_ocr
            return texto
        
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            return await self.extraer_texto_imagen(contenido, archivo.filename)
        
        elif extension in ['.xlsx', '.xls']:
            return await self.extraer_texto_excel(contenido, archivo.filename)
        
        elif extension in ['.docx', '.doc']:
            return await self.extraer_texto_word(contenido, archivo.filename)
        
        else:
            raise ValueError(f"Tipo de archivo no soportado: {extension}")
    
    async def extraer_texto_pdf(self, contenido: bytes, nombre_archivo: str = "documento.pdf") -> str:
        """
        Extrae texto de archivo PDF usando PyPDF2.
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario del archivo PDF
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído del PDF
        """
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(contenido))
            texto_completo = ""
            
            num_paginas = len(pdf_reader.pages)
            logger.info(f"📖 Procesando PDF con {num_paginas} página(s)")
            
            for i, page in enumerate(pdf_reader.pages):
                texto_pagina = page.extract_text()
                texto_completo += f"\n--- PÁGINA {i+1} ---\n{texto_pagina}\n"
            
            texto_final = texto_completo.strip()
            
            # Preparar metadatos
            metadatos = {
                "total_paginas": num_paginas,
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "PyPDF2",
                "caracteres_extraidos": len(texto_final)
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_final, "PDF", metadatos
            )
            
            logger.info(f"✅ Texto extraído de PDF: {len(texto_final)} caracteres")
            
            return texto_final
            
        except Exception as e:
            error_msg = f"Error procesando PDF: {str(e)}"
            logger.error(f"❌ {error_msg}")
            
            # Guardar también los errores para debugging
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "PDF_ERROR", {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_pdf_con_ocr(self, contenido_pdf: bytes, nombre_archivo: str = "documento.pdf") -> str:
        """
        Extrae texto de PDF convirtiéndolo a imágenes y aplicando OCR.
        NUEVA FUNCIÓN para manejar PDFs con poco texto extraíble.
        
        Args:
            contenido_pdf: Contenido binario del archivo PDF
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído via OCR de las imágenes del PDF
        """
        try:
            # VALIDAR PDF ANTES DE INTENTAR CONVERSIÓN
            validacion = self._validar_pdf(contenido_pdf, nombre_archivo)
            
            if not validacion["valido"]:
                error_msg = f"PDF no válido para OCR: {validacion['error']}"
                logger.error(f"❌ {error_msg}")
                
                # Guardar error detallado
                self._guardar_texto_extraido(
                    nombre_archivo, error_msg, "PDF_OCR_ERROR", 
                    {"error": validacion["error"], "validacion": validacion}
                )
                
                return error_msg
            
            logger.info(f"✅ PDF validado: {validacion['info']['paginas']} páginas, {validacion['info']['tamaño_bytes']} bytes")
            logger.info(f"🔧 Método recomendado: {validacion['metodo_recomendado']}")
            
            # Convertir PDF a imágenes
            imagenes = self._convertir_pdf_a_imagenes(contenido_pdf, nombre_archivo)
            
            if not imagenes:
                error_msg = "No se pudieron convertir páginas del PDF a imágenes"
                logger.error(f"❌ {error_msg}")
                
                self._guardar_texto_extraido(
                    nombre_archivo, error_msg, "PDF_OCR_ERROR", 
                    {"error": "Conversión PDF → Imagen falló", "validacion": validacion}
                )
                
                return error_msg
            
            # Aplicar OCR a cada imagen
            texto_total = ""
            total_caracteres = 0
            
            for i, imagen_bytes in enumerate(imagenes):
                logger.info(f"🔍 Aplicando OCR a página {i+1}/{len(imagenes)}")
                
                # Aplicar OCR a esta imagen
                texto_pagina = await self._aplicar_ocr_a_imagen(imagen_bytes, f"página_{i+1}")
                
                if texto_pagina and not texto_pagina.startswith("Error"):
                    texto_total += f"\n--- PÁGINA {i+1} (OCR) ---\n{texto_pagina}\n"
                    total_caracteres += len(texto_pagina)
                else:
                    texto_total += f"\n--- PÁGINA {i+1} (OCR) ---\n[Error en OCR o página vacía]\n"
            
            # Preparar metadatos
            metadatos = {
                "total_paginas": len(imagenes),
                "tamaño_archivo_bytes": len(contenido_pdf),
                "metodo": "PDF → Imagen → Google Vision OCR",
                "caracteres_extraidos": total_caracteres,
                "paginas_procesadas": len(imagenes),
                "validacion": validacion["info"]
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_total, "PDF_OCR", metadatos
            )
            
            logger.info(f"✅ OCR de PDF completado: {total_caracteres} caracteres de {len(imagenes)} páginas")
            
            return texto_total.strip()
            
        except Exception as e:
            error_msg = f"Error en OCR de PDF: {str(e)}"
            logger.error(f"❌ {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "PDF_OCR_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def _aplicar_ocr_a_imagen(self, imagen_bytes: bytes, descripcion: str = "imagen") -> str:
        """
        Aplica OCR a una imagen específica.
        
        Args:
            imagen_bytes: Bytes de la imagen
            descripcion: Descripción para logging
            
        Returns:
            str: Texto extraído de la imagen
        """
        if not self.vision_client:
            return "OCR no disponible - Google Vision no configurado"
        
        try:
            # Crear objeto Image para Vision
            image = vision.Image(content=imagen_bytes)
            
            # Detectar texto
            response = self.vision_client.text_detection(image=image)
            
            if response.error.message:
                raise Exception(f'Error en Vision API: {response.error.message}')
            
            texts = response.text_annotations
            
            if texts:
                texto_extraido = texts[0].description
                logger.debug(f"✅ OCR exitoso en {descripcion}: {len(texto_extraido)} caracteres")
                return texto_extraido
            else:
                logger.debug(f"⚠️ No se detectó texto en {descripcion}")
                return ""
                
        except Exception as e:
            logger.error(f"❌ Error en OCR de {descripcion}: {e}")
            return f"Error en OCR: {str(e)}"
    
    async def extraer_texto_imagen(self, contenido: bytes, nombre_archivo: str = "imagen.jpg", metodo: str = "OCR") -> str:
        """
        Extrae texto de imagen usando Google Vision OCR.
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario de la imagen
            nombre_archivo: Nombre del archivo original para guardado
            metodo: Método específico de OCR (para casos como fallback)
            
        Returns:
            str: Texto extraído de la imagen
        """
        if not self.vision_client:
            error_msg = "OCR no disponible - Google Vision no configurado"
            logger.warning(f"⚠️ {error_msg}")
            
            # Guardar el error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, f"{metodo}_ERROR", 
                {"error": "Google Vision no configurado"}
            )
            
            return error_msg
        
        try:
            # Aplicar OCR
            texto_extraido = await self._aplicar_ocr_a_imagen(contenido, nombre_archivo)
            
            if texto_extraido and not texto_extraido.startswith("Error"):
                # Preparar metadatos
                metadatos = {
                    "tamaño_archivo_bytes": len(contenido),
                    "metodo": "Google Vision OCR",
                    "caracteres_extraidos": len(texto_extraido)
                }
                
                # Guardar texto extraído automáticamente
                archivo_guardado = self._guardar_texto_extraido(
                    nombre_archivo, texto_extraido, metodo, metadatos
                )
                
                logger.info(f"✅ OCR exitoso: {len(texto_extraido)} caracteres extraídos")
                return texto_extraido
            else:
                no_texto_msg = "No se detectó texto en la imagen"
                logger.warning(f"⚠️ {no_texto_msg}")
                
                # Guardar resultado vacío
                self._guardar_texto_extraido(
                    nombre_archivo, no_texto_msg, f"{metodo}_VACIO", 
                    {"elementos_detectados": 0}
                )
                
                return no_texto_msg
                
        except Exception as e:
            error_msg = f"Error en OCR con Google Vision: {str(e)}"
            logger.error(f"❌ {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, f"{metodo}_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_excel(self, contenido: bytes, nombre_archivo: str = "archivo.xlsx") -> str:
        """
        Extrae texto de archivo Excel (XLSX/XLS).
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario del archivo Excel
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído y formateado del Excel
        """
        try:
            # Leer Excel
            df = pd.read_excel(io.BytesIO(contenido), sheet_name=None)  # Leer todas las hojas
            
            texto_completo = ""
            total_hojas = 0
            total_filas = 0
            
            # Si hay múltiples hojas
            if isinstance(df, dict):
                total_hojas = len(df)
                for nombre_hoja, dataframe in df.items():
                    texto_completo += f"\n--- HOJA: {nombre_hoja} ---\n"
                    texto_completo += dataframe.to_string(index=False, na_rep='')
                    texto_completo += "\n"
                    total_filas += len(dataframe)
            else:
                # Una sola hoja
                total_hojas = 1
                total_filas = len(df)
                texto_completo = df.to_string(index=False, na_rep='')
            
            texto_final = texto_completo.strip()
            
            # Preparar metadatos
            metadatos = {
                "total_hojas": total_hojas,
                "total_filas": total_filas,
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "pandas.read_excel"
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_final, "EXCEL", metadatos
            )
            
            logger.info(f"✅ Excel procesado: {len(texto_final)} caracteres extraídos")
            logger.info(f"📊 Hojas: {total_hojas}, Filas: {total_filas}")
            
            return texto_final
            
        except Exception as e:
            error_msg = f"Error procesando Excel: {str(e)}"
            logger.error(f"❌ {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "EXCEL_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    async def extraer_texto_word(self, contenido: bytes, nombre_archivo: str = "documento.docx") -> str:
        """
        Extrae texto de archivo Word (DOCX).
        GUARDA AUTOMÁTICAMENTE el texto extraído.
        
        Args:
            contenido: Contenido binario del archivo Word
            nombre_archivo: Nombre del archivo original para guardado
            
        Returns:
            str: Texto extraído del documento Word
        """
        try:
            # Leer documento Word
            doc = Document(io.BytesIO(contenido))
            texto_completo = ""
            total_parrafos = 0
            total_tablas = 0
            
            # Extraer párrafos
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():  # Solo párrafos con contenido
                    texto_completo += paragraph.text + "\n"
                    total_parrafos += 1
            
            # Extraer tablas si las hay
            if doc.tables:
                total_tablas = len(doc.tables)
                texto_completo += "\n--- TABLAS DEL DOCUMENTO ---\n"
                for i, table in enumerate(doc.tables):
                    texto_completo += f"\nTabla {i+1}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        if row_text.strip():
                            texto_completo += row_text + "\n"
            
            texto_final = texto_completo.strip()
            
            # Preparar metadatos
            metadatos = {
                "total_parrafos": total_parrafos,
                "total_tablas": total_tablas,
                "tamaño_archivo_bytes": len(contenido),
                "metodo": "python-docx"
            }
            
            # Guardar texto extraído automáticamente
            archivo_guardado = self._guardar_texto_extraido(
                nombre_archivo, texto_final, "WORD", metadatos
            )
            
            logger.info(f"✅ Word procesado: {len(texto_final)} caracteres extraídos")
            logger.info(f"📄 Párrafos: {total_parrafos}, Tablas: {total_tablas}")
            
            return texto_final
            
        except Exception as e:
            error_msg = f"Error procesando Word: {str(e)}"
            logger.error(f"❌ {error_msg}")
            
            # Guardar error
            self._guardar_texto_extraido(
                nombre_archivo, error_msg, "WORD_ERROR", 
                {"error": str(e)}
            )
            
            return error_msg
    
    def validar_archivo(self, archivo: UploadFile) -> Dict[str, Any]:
        """
        Valida si un archivo es procesable y retorna información sobre él.
        
        Args:
            archivo: Archivo a validar
            
        Returns:
            Dict con información de validación
        """
        extensiones_soportadas = ['.pdf', '.jpg', '.jpeg', '.png', '.gif', '.bmp', 
                                '.tiff', '.xlsx', '.xls', '.docx', '.doc']
        
        if not archivo.filename:
            return {
                "valido": False,
                "error": "Archivo sin nombre"
            }
        
        extension = Path(archivo.filename).suffix.lower()
        
        if extension not in extensiones_soportadas:
            return {
                "valido": False,
                "error": f"Extensión no soportada: {extension}",
                "extensiones_soportadas": extensiones_soportadas
            }
        
        # Determinar tipo de procesamiento
        tipo_procesamiento = "Desconocido"
        guardado_automatico = True
        
        if extension == '.pdf':
            if PDF2IMAGE_DISPONIBLE or PYMUPDF_DISPONIBLE:
                tipo_procesamiento = "Extracción PDF + OCR con conversión a imagen"
            else:
                tipo_procesamiento = "Extracción PDF (OCR fallback limitado)"
        elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff']:
            tipo_procesamiento = "OCR con Google Vision"
        elif extension in ['.xlsx', '.xls']:
            tipo_procesamiento = "Procesamiento Excel"
        elif extension in ['.docx', '.doc']:
            tipo_procesamiento = "Procesamiento Word"
        
        return {
            "valido": True,
            "extension": extension,
            "tipo_procesamiento": tipo_procesamiento,
            "ocr_disponible": self.vision_client is not None,
            "pdf_to_image_disponible": PDF2IMAGE_DISPONIBLE or PYMUPDF_DISPONIBLE,
            "guardado_automatico": guardado_automatico,
            "carpeta_guardado": str(self.carpeta_fecha)
        }
    
    async def procesar_multiples_archivos(self, archivos: list) -> Dict[str, str]:
        """
        Procesa múltiples archivos y retorna diccionario con textos extraídos.
        GUARDA AUTOMÁTICAMENTE cada extracción individual.
        
        Args:
            archivos: Lista de archivos UploadFile
            
        Returns:
            Dict[str, str]: Diccionario {nombre_archivo: texto_extraido}
        """
        textos_extraidos = {}
        
        logger.info(f"📁 Procesando {len(archivos)} archivos con guardado automático")
        logger.info(f"💾 Carpeta de guardado: {self.carpeta_fecha}")
        
        for archivo in archivos:
            try:
                # Validar archivo
                validacion = self.validar_archivo(archivo)
                if not validacion["valido"]:
                    logger.error(f"❌ Archivo inválido {archivo.filename}: {validacion['error']}")
                    textos_extraidos[archivo.filename] = f"ERROR: {validacion['error']}"
                    continue
                
                # Procesar archivo (esto automáticamente guarda el texto)
                texto = await self.procesar_archivo(archivo)
                textos_extraidos[archivo.filename] = texto
                
                logger.info(f"✅ Archivo procesado y guardado: {archivo.filename}")
                
            except Exception as e:
                logger.error(f"❌ Error procesando archivo {archivo.filename}: {e}")
                textos_extraidos[archivo.filename] = f"ERROR PROCESANDO: {str(e)}"
                
                # Guardar también los errores de procesamiento
                self._guardar_texto_extraido(
                    archivo.filename, f"ERROR PROCESANDO: {str(e)}", 
                    "PROCESAMIENTO_ERROR", {"error": str(e)}
                )
        
        logger.info(f"🎉 Procesamiento completado: {len(textos_extraidos)} archivos")
        logger.info(f"📂 Todos los textos extraídos guardados en: {self.carpeta_fecha}")
        
        return textos_extraidos
    
    def obtener_estadisticas_guardado(self) -> Dict[str, Any]:
        """
        Obtiene estadísticas de los archivos guardados hoy.
        
        Returns:
            Dict con estadísticas de guardado
        """
        try:
            archivos_guardados = list(self.carpeta_fecha.glob("*.txt"))
            
            estadisticas = {
                "fecha": datetime.now().strftime("%Y-%m-%d"),
                "carpeta": str(self.carpeta_fecha),
                "total_archivos_guardados": len(archivos_guardados),
                "tipos_extraccion": {},
                "tamaño_total_mb": 0,
                "dependencias": {
                    "google_vision": self.vision_client is not None,
                    "pdf2image": PDF2IMAGE_DISPONIBLE,
                    "pymupdf": PYMUPDF_DISPONIBLE
                }
            }
            
            # Analizar archivos guardados
            for archivo in archivos_guardados:
                # Extraer tipo de extracción del nombre
                partes = archivo.name.split("_")
                if len(partes) >= 2:
                    tipo = partes[1]
                    estadisticas["tipos_extraccion"][tipo] = estadisticas["tipos_extraccion"].get(tipo, 0) + 1
                
                # Sumar tamaño
                estadisticas["tamaño_total_mb"] += archivo.stat().st_size / (1024 * 1024)
            
            estadisticas["tamaño_total_mb"] = round(estadisticas["tamaño_total_mb"], 2)
            
            return estadisticas
            
        except Exception as e:
            logger.error(f"Error obteniendo estadísticas: {e}")
            return {"error": str(e)}
